import io
import os
import sys
import traceback
import logging
import requests
# Используем новый SDK, как просит терминал
from google import genai
from google.genai import types
from vkbottle import Keyboard, KeyboardButtonColor, Text, DocMessagesUploader, PhotoMessageUploader
from vkbottle.bot import Bot, Message
from docx import Document
from dotenv import load_dotenv
from typing import Dict

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
    stream=sys.stdout,
)


def mask_secret(value: str | None) -> str:
    if not value:
        return "missing"
    if len(value) <= 8:
        return "***"
    return f"{value[:4]}...{value[-4:]}"


# ================= НАСТРОЙКИ =================
load_dotenv()
VK_TOKEN = os.getenv("VK_TOKEN")
GEMINI_KEY = os.getenv("GEMINI_KEY")
HF_TOKEN = os.getenv("HF_TOKEN")

required_vars = {
    "VK_TOKEN": VK_TOKEN,
    "GEMINI_KEY": GEMINI_KEY,
    "HF_TOKEN": HF_TOKEN,
}
missing_vars = [name for name, value in required_vars.items() if not value]
if missing_vars:
    logging.error("Missing environment variables: %s", ", ".join(missing_vars))
    logging.error("Expected vars: VK_TOKEN, GEMINI_KEY, HF_TOKEN")
    logging.error("Current VK_TOKEN: %s", mask_secret(VK_TOKEN))
    logging.error("Current GEMINI_KEY: %s", mask_secret(GEMINI_KEY))
    logging.error("Current HF_TOKEN: %s", mask_secret(HF_TOKEN))
    raise RuntimeError("Environment variables are missing")

# Инициализируем нового клиента Gemini
client = genai.Client(api_key=GEMINI_KEY)
MODEL_NAME = "gemini-2.5-flash" 

bot = Bot(token=VK_TOKEN)
photo_uploader = PhotoMessageUploader(bot.api)
doc_uploader = DocMessagesUploader(bot.api)

sessions: Dict[int, dict] = {}

SYSTEM_PROMPT = (
    "ты гений с мировым именем но при этом дерзкий и хамоватый тип "
    "на научные темы отвечай академично и сложно как профессор "
    "в бытовом общении пиши только с маленькой буквы без запятых юзай сленг типа крутой челик на чилле "
    "импровизируй в хамстве не повторяйся"
)

def get_format_keyboard():
    kb = Keyboard(one_time=True, inline=False)
    kb.add(Text("скинь текстом"), color=KeyboardButtonColor.PRIMARY)
    kb.add(Text("сделай файл"), color=KeyboardButtonColor.POSITIVE)
    return kb.get_json()

@bot.on.message(text="завершить диалог")
async def finish_handler(message: Message):
    if message.from_id in sessions:
        del sessions[message.from_id]
    await message.answer("все замяли тему я все забыл")

@bot.on.message(text="скинь текстом")
async def send_as_text(message: Message):
    uid = message.from_id
    if uid in sessions and sessions[uid].get("last_res"):
        res = sessions[uid]["last_res"]
        for i in range(0, len(res), 4000):
            await message.answer(res[i:i+4000])
    else:
        await message.answer("че ты жмешь я еще не родил ответ")

@bot.on.message(text="сделай файл")
async def send_as_file(message: Message):
    uid = message.from_id
    if uid in sessions and sessions[uid].get("last_res"):
        doc_stream = io.BytesIO()
        doc = Document()
        doc.add_paragraph(sessions[uid]["last_res"])
        doc.save(doc_stream)
        doc_stream.seek(0)
        uploaded = await doc_uploader.upload(doc_stream, peer_id=uid, title="answer.docx")
        await message.answer("на держи свой файл", attachment=uploaded)

@bot.on.message()
async def main_handler(message: Message):
    uid = message.from_id
    text = message.text or ""

    if text.lower() in ["скинь текстом", "сделай файл"]:
        return

    # Обработка вложений
    prompt_parts = [text] if text else []
    
    if message.attachments:
        await message.answer("вижу твой хлам ща чекну...")
        for att in message.attachments:
            if att.photo:
                photo_url = att.photo.sizes[-1].url
                photo_data = requests.get(photo_url).content
                prompt_parts.append(types.Part.from_bytes(data=photo_data, mime_type="image/jpeg"))
            elif att.doc and att.doc.ext == "docx":
                doc_data = requests.get(att.doc.url).content
                doc_obj = Document(io.BytesIO(doc_data))
                full_text = "\n".join([p.text for p in doc_obj.paragraphs])
                prompt_parts.append(f"\n[Файл {att.doc.title}]:\n{full_text}")

    if text.lower().startswith("нарисуй"):
        query = text.lower().replace("нарисуй", "").strip()
        await message.answer("рисую не зуди...")
        try:
            res = requests.post("https://api-inference.huggingface.co/models/black-forest-labs/FLUX.1-schnell", 
                                headers={"Authorization": f"Bearer {HF_TOKEN}"}, json={"inputs": query})
            photo = await photo_uploader.upload(io.BytesIO(res.content), peer_id=uid)
            await message.answer(f"на: {query}", attachment=photo)
        except:
            await message.answer("краски кончились")
        return

    await message.answer("погоди рожаю ответ...")
    
    try:
        # Новая логика вызова API
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=prompt_parts,
            config=types.GenerateContentConfig(
                system_instruction=SYSTEM_PROMPT,
                safety_settings=[types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE")]
            )
        )
        
        ans_text = response.text
        sessions[uid] = {"last_res": ans_text}
        
        await message.answer(ans_text)
        await message.answer("как скинуть это?", keyboard=get_format_keyboard())
            
    except Exception as e:
        await message.answer(f"мозг поплыл: {e}")

if __name__ == "__main__":
    try:
        logging.info("Booting GeminiVKBOT container process")
        logging.info("VK_TOKEN: %s", mask_secret(VK_TOKEN))
        logging.info("GEMINI_KEY: %s", mask_secret(GEMINI_KEY))
        logging.info("HF_TOKEN: %s", mask_secret(HF_TOKEN))
        logging.info("--- БОТ ЗАПУЩЕН И ГОТОВ ХАМИТЬ ---")
        bot.run_forever()
    except Exception:
        logging.error("Fatal startup/runtime error in bot process")
        traceback.print_exc()
        raise
